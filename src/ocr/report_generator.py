import csv
import json
import logging
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, List

import pandas as pd
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt

class ReportGenerator:
    """A class to generate formatted reports from extracted document data."""
    
    def __init__(self, output_dir: str):
        """
        Initialize the ReportGenerator.
        
        Args:
            output_dir (str): Directory to save generated reports
        """
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.logger = self._setup_logger()
    
    def _setup_logger(self):
        """Set up logging configuration."""
        logger = logging.getLogger('ReportGenerator')
        logger.setLevel(logging.INFO)
        
        if not logger.handlers:
            handler = logging.StreamHandler()
            formatter = logging.Formatter(
                '%(asctime)s - %(name)s - %(levelname)s - %(message)s'
            )
            handler.setFormatter(formatter)
            logger.addHandler(handler)
            
        return logger
    
    def generate_csv(self, data: Dict[str, Any], filename: str) -> str:
        """
        Generate a CSV report from extracted data.
        
        Args:
            data (Dict[str, Any]): Extracted data
            filename (str): Output filename
            
        Returns:
            str: Path to generated CSV file
        """
        try:
            output_path = self.output_dir / f"{filename}.csv"
            
            # Flatten the data structure
            flattened_data = []
            for key, values in data.items():
                if isinstance(values, list) and key != 'tables':
                    for value in values:
                        flattened_data.append({
                            'type': key,
                            'value': value,
                            'timestamp': data.get('timestamp', '')
                        })
            
            # Write to CSV
            if flattened_data:
                df = pd.DataFrame(flattened_data)
                df.to_csv(output_path, index=False)
                self.logger.info(f"CSV report generated: {output_path}")
                return str(output_path)
            
        except Exception as e:
            self.logger.error(f"Error generating CSV report: {str(e)}")
            raise
    
    def generate_docx(self, data: Dict[str, Any], filename: str) -> str:
        """
        Generate a Word document report from extracted data.
        
        Args:
            data (Dict[str, Any]): Extracted data
            filename (str): Output filename
            
        Returns:
            str: Path to generated Word document
        """
        try:
            output_path = self.output_dir / f"{filename}.docx"
            
            doc = Document()
            doc.add_heading('Document Analysis Report', 0)
            
            # Add timestamp
            doc.add_paragraph(f"Generated on: {data.get('timestamp', datetime.now().isoformat())}")
            
            # Add sections for each data type
            for key, values in data.items():
                if isinstance(values, list) and key != 'tables' and key != 'raw_text':
                    if values:
                        doc.add_heading(key.replace('_', ' ').title(), level=1)
                        for value in values:
                            doc.add_paragraph(f"• {value}")
            
            # Add tables if present
            if data.get('tables'):
                doc.add_heading('Extracted Tables', level=1)
                for table_data in data['tables']:
                    table = doc.add_table(rows=1, cols=len(table_data))
                    table.style = 'Table Grid'
                    
                    for i, value in enumerate(table_data):
                        table.cell(0, i).text = str(value)
            
            # Add raw text at the end
            if data.get('raw_text'):
                doc.add_heading('Raw Extracted Text', level=1)
                doc.add_paragraph(data['raw_text'])
            
            doc.save(output_path)
            self.logger.info(f"Word document report generated: {output_path}")
            return str(output_path)
            
        except Exception as e:
            self.logger.error(f"Error generating Word document: {str(e)}")
            raise
    
    def generate_summary_visualizations(self, data: Dict[str, Any], filename: str) -> str:
        """
        Generate summary visualizations from extracted data.
        
        Args:
            data (Dict[str, Any]): Extracted data
            filename (str): Output filename
            
        Returns:
            str: Path to generated visualization file
        """
        try:
            output_path = self.output_dir / f"{filename}_summary.png"
            
            # Create a summary of counts
            counts = {
                key: len(values) for key, values in data.items()
                if isinstance(values, list) and key not in ['tables', 'raw_text']
            }
            
            if counts:
                # Create bar plot
                plt.figure(figsize=(10, 6))
                plt.bar(counts.keys(), counts.values())
                plt.xticks(rotation=45)
                plt.title('Document Content Summary')
                plt.ylabel('Count')
                plt.tight_layout()
                
                # Save plot
                plt.savefig(output_path)
                plt.close()
                
                self.logger.info(f"Summary visualization generated: {output_path}")
                return str(output_path)
            
        except Exception as e:
            self.logger.error(f"Error generating visualization: {str(e)}")
            raise
    
    def generate_reports(self, data: Dict[str, Any], base_filename: str) -> Dict[str, str]:
        """
        Generate all report formats.
        
        Args:
            data (Dict[str, Any]): Extracted data
            base_filename (str): Base name for output files
            
        Returns:
            Dict[str, str]: Paths to generated reports
        """
        try:
            self.logger.info("Starting report generation")
            
            reports = {}
            
            # Generate CSV report
            csv_path = self.generate_csv(data, base_filename)
            if csv_path:
                reports['csv'] = csv_path
            
            # Generate Word document
            docx_path = self.generate_docx(data, base_filename)
            if docx_path:
                reports['docx'] = docx_path
            
            # Generate visualizations
            viz_path = self.generate_summary_visualizations(data, base_filename)
            if viz_path:
                reports['visualization'] = viz_path
            
            # Save raw data as JSON
            json_path = self.output_dir / f"{base_filename}.json"
            with open(json_path, 'w') as f:
                json.dump(data, f, indent=2)
            reports['json'] = str(json_path)
            
            self.logger.info("Report generation completed successfully")
            return reports
            
        except Exception as e:
            self.logger.error(f"Error generating reports: {str(e)}")
            raise